"""
Bygger rapport.docx fra alle markdown-filer i 005_report/.
Kjør fra prosjektrot: python 004_scripts/bygg_rapport_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = Path(__file__).parent.parent
REPORT = ROOT / "005_report"
OUT    = ROOT / "rapport.docx"

# Rekkefølge på filer
FILER = [
    "00_sammendrag.md",
    "01_introduksjon.md",
    "02_teori_og_litteratur.md",
    "03_casebeskrivelse.md",
    "04_data_og_metode.md",
    "05_modellering.md",
    "06_analyse_og_resultater.md",
    "07_diskusjon.md",
    "08_konklusjon.md",
]

REFERANSER = ROOT / "003_references" / "referanser.md"


# ---------------------------------------------------------------------------
# Hjelpefunksjoner for formatering
# ---------------------------------------------------------------------------

def sett_font(run, navn="Times New Roman", størrelse=12, fet=False,
              kursiv=False, farge=None):
    run.font.name = navn
    run.font.size = Pt(størrelse)
    run.font.bold = fet
    run.font.italic = kursiv
    if farge:
        run.font.color.rgb = RGBColor(*farge)


def sett_avsnitt_format(para, venstre=0, første=0,
                        mellom_avsnitt=6, linjeavstand=None):
    pf = para.paragraph_format
    pf.left_indent      = Cm(venstre)
    pf.first_line_indent = Cm(første)
    pf.space_after      = Pt(mellom_avsnitt)
    if linjeavstand:
        from docx.shared import Length
        pf.line_spacing = linjeavstand


def legg_til_sideskift(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


def legg_til_horisontal_linje(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)


# ---------------------------------------------------------------------------
# Parser for én markdown-linje
# ---------------------------------------------------------------------------

def parse_inline(tekst: str) -> list[tuple[str, bool, bool, bool]]:
    """
    Returnerer liste av (tekst, fet, kursiv, kode).
    Støtter **fet**, *kursiv*, `kode` og kombinasjoner.
    """
    mønster = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    deler   = []
    sist    = 0
    for m in mønster.finditer(tekst):
        if m.start() > sist:
            deler.append((tekst[sist:m.start()], False, False, False))
        if m.group(2):
            deler.append((m.group(2), True, True, False))
        elif m.group(3):
            deler.append((m.group(3), True, False, False))
        elif m.group(4):
            deler.append((m.group(4), False, True, False))
        elif m.group(5):
            deler.append((m.group(5), False, False, True))
        sist = m.end()
    if sist < len(tekst):
        deler.append((tekst[sist:], False, False, False))
    return deler or [(tekst, False, False, False)]


def legg_til_løpetekst(para, tekst: str, grunnstørrelse=12):
    for del_tekst, fet, kursiv, kode in parse_inline(tekst):
        run = para.add_run(del_tekst)
        if kode:
            sett_font(run, navn="Courier New", størrelse=10)
        else:
            sett_font(run, størrelse=grunnstørrelse, fet=fet, kursiv=kursiv)


# ---------------------------------------------------------------------------
# Hoved-parser: markdown → docx
# ---------------------------------------------------------------------------

def parse_markdown(doc: Document, md_tekst: str, er_sammendrag=False):
    linjer      = md_tekst.splitlines()
    i           = 0
    i_kodeblokk = False
    kode_buffer = []

    while i < len(linjer):
        linje = linjer[i]

        # --- Kodeblokk ---
        if linje.strip().startswith("```"):
            if not i_kodeblokk:
                i_kodeblokk = True
                kode_buffer = []
            else:
                # Skriv ut kodeblokken
                para = doc.add_paragraph()
                para.paragraph_format.left_indent  = Cm(1)
                para.paragraph_format.space_after  = Pt(6)
                para.paragraph_format.space_before = Pt(6)
                run  = para.add_run("\n".join(kode_buffer))
                sett_font(run, navn="Courier New", størrelse=9)
                # Lys grå bakgrunn via shading
                pPr   = para._p.get_or_add_pPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "F2F2F2")
                pPr.append(shd)
                i_kodeblokk = False
                kode_buffer = []
            i += 1
            continue

        if i_kodeblokk:
            kode_buffer.append(linje)
            i += 1
            continue

        # --- Hopp over skrivenotater og HTML-kommentarer ---
        if linje.strip().startswith("<!--") or "Skrivenotater" in linje:
            i += 1
            continue

        # --- Figurreferanse → sett inn bilde ---
        figur_match = re.search(r"→\s*Figur[:\s]+`?([^`\s]+\.png)`?", linje)
        if figur_match:
            # Støtter både bare filnavn og full sti
            sti_str  = figur_match.group(1)
            filnavn  = Path(sti_str).name
            bildefil = REPORT / "figures" / filnavn
            if bildefil.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run  = para.add_run()
                run.add_picture(str(bildefil), width=Cm(14))
                # Figurtekst under bildet
                cap_tekst = linje.strip().lstrip("→").strip()
                cap_tekst = re.sub(r"\S+\.png", "", cap_tekst).strip(": ").strip()
                if cap_tekst:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r   = cap.add_run(cap_tekst)
                    sett_font(r, størrelse=10, kursiv=True)
                    cap.paragraph_format.space_after = Pt(10)
            else:
                # Fil ikke funnet – behold som tekst
                para = doc.add_paragraph()
                legg_til_løpetekst(para, linje.strip())
            i += 1
            continue

        # --- Tom linje ---
        if not linje.strip():
            i += 1
            continue

        # --- Horisontal linje ---
        if linje.strip() in ("---", "***", "___"):
            legg_til_horisontal_linje(doc)
            i += 1
            continue

        # --- Overskrifter ---
        if linje.startswith("# "):
            tekst = linje[2:].strip()
            if er_sammendrag:
                para = doc.add_paragraph()
                run  = para.add_run(tekst)
                sett_font(run, størrelse=16, fet=True)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(12)
            # Kapitteloverskrift håndteres av kallstedet
            i += 1
            continue

        if linje.startswith("## "):
            tekst = linje[3:].strip()
            para  = doc.add_heading(tekst, level=2)
            for run in para.runs:
                sett_font(run, størrelse=13, fet=True)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        if linje.startswith("### "):
            tekst = linje[4:].strip()
            para  = doc.add_heading(tekst, level=3)
            for run in para.runs:
                sett_font(run, størrelse=12, fet=True)
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # --- Sitatblokk (blockquote) ---
        if linje.startswith("> "):
            tekst = linje[2:].strip()
            # Fjern **bold** fra problemstillingen
            tekst = re.sub(r"\*\*(.+?)\*\*", r"\1", tekst)
            para  = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(1.5)
            para.paragraph_format.space_after  = Pt(6)
            run   = para.add_run(tekst)
            sett_font(run, størrelse=12, kursiv=True)
            i += 1
            continue

        # --- Tabell ---
        if linje.startswith("|"):
            tabell_linjer = []
            while i < len(linjer) and linjer[i].startswith("|"):
                if not re.match(r"^\|[-| :]+\|$", linjer[i].strip()):
                    tabell_linjer.append(linjer[i])
                i += 1
            if not tabell_linjer:
                continue

            rader = []
            for tl in tabell_linjer:
                celler = [c.strip() for c in tl.strip().strip("|").split("|")]
                rader.append(celler)

            if not rader:
                continue

            n_kol  = len(rader[0])
            tabell = doc.add_table(rows=len(rader), cols=n_kol)
            tabell.style = "Table Grid"

            for r_idx, rad in enumerate(rader):
                for k_idx, celle_tekst in enumerate(rad):
                    celle = tabell.cell(r_idx, k_idx)
                    celle.text = ""
                    para  = celle.paragraphs[0]
                    # Fjern markdown-formatering fra tabellceller
                    ren   = re.sub(r"\*\*(.+?)\*\*", r"\1", celle_tekst)
                    ren   = re.sub(r"\*(.+?)\*",     r"\1", ren)
                    run   = para.add_run(ren)
                    er_header = (r_idx == 0)
                    sett_font(run, størrelse=10, fet=er_header)
                    para.paragraph_format.space_after  = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    if er_header:
                        # Grå bakgrunn på header-rad
                        tc_pr = celle._tc.get_or_add_tcPr()
                        shd   = OxmlElement("w:shd")
                        shd.set(qn("w:val"),   "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"),  "D9D9D9")
                        tc_pr.append(shd)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # --- Punktliste ---
        if linje.startswith("- ") or linje.startswith("* "):
            tekst = linje[2:].strip()
            para  = doc.add_paragraph(style="List Bullet")
            legg_til_løpetekst(para, tekst)
            para.paragraph_format.left_indent   = Cm(0.75)
            para.paragraph_format.space_after   = Pt(2)
            i += 1
            continue

        # --- Nummerert liste ---
        if re.match(r"^\d+\.\s", linje):
            tekst = re.sub(r"^\d+\.\s", "", linje)
            para  = doc.add_paragraph(style="List Number")
            legg_til_løpetekst(para, tekst)
            para.paragraph_format.left_indent  = Cm(0.75)
            para.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # --- Kursiv-/kildehenvisning (linje starter med *) ---
        if linje.startswith("*") and linje.endswith("*") and not linje.startswith("**"):
            tekst = linje.strip("*")
            para  = doc.add_paragraph()
            run   = para.add_run(tekst)
            sett_font(run, størrelse=10, kursiv=True)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # --- Vanlig avsnitt ---
        para = doc.add_paragraph()
        legg_til_løpetekst(para, linje.strip())
        para.paragraph_format.space_after = Pt(6)
        i += 1


# ---------------------------------------------------------------------------
# Forside
# ---------------------------------------------------------------------------

def lag_forside(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Kurs
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LOG650 – Kvantitative metoder i logistikk")
    sett_font(r, størrelse=13)

    doc.add_paragraph()

    # Tittel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Etterspørselsprognoser og KI-støttet lagerstyring")
    sett_font(r, størrelse=20, fet=True)

    doc.add_paragraph()

    # Undertittel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Byggmakker Gravdal")
    sett_font(r, størrelse=14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Forfatter
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Oskar Eide")
    sett_font(r, størrelse=13, fet=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Høgskolen i Molde")
    sett_font(r, størrelse=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026")
    sett_font(r, størrelse=12)

    legg_til_sideskift(doc)


# ---------------------------------------------------------------------------
# Hoved
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Sideformat: A4
    for seksjon in doc.sections:
        seksjon.page_width  = Cm(21)
        seksjon.page_height = Cm(29.7)
        seksjon.left_margin   = Cm(3.5)
        seksjon.right_margin  = Cm(2.5)
        seksjon.top_margin    = Cm(2.5)
        seksjon.bottom_margin = Cm(2.5)

    # Standard avsnittsstil
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)

    # --- Forside ---
    lag_forside(doc)

    # --- Sammendrag ---
    print("Behandler: 00_sammendrag.md")
    tekst = (REPORT / "00_sammendrag.md").read_text(encoding="utf-8")
    parse_markdown(doc, tekst, er_sammendrag=True)
    legg_til_sideskift(doc)

    # --- Kapittel 1–8 ---
    kapittel_nr = {
        "01_introduksjon.md":          ("1", "Introduksjon"),
        "02_teori_og_litteratur.md":   ("2", "Teori og litteratur"),
        "03_casebeskrivelse.md":       ("3", "Casebeskrivelse"),
        "04_data_og_metode.md":        ("4", "Data og metode"),
        "05_modellering.md":           ("5", "Modellering"),
        "06_analyse_og_resultater.md": ("6", "Analyse og resultater"),
        "07_diskusjon.md":             ("7", "Diskusjon"),
        "08_konklusjon.md":            ("8", "Konklusjon"),
    }

    for filnavn in FILER[1:]:
        sti = REPORT / filnavn
        if not sti.exists():
            print(f"  Hopper over (ikke funnet): {filnavn}")
            continue

        print(f"Behandler: {filnavn}")
        tekst  = sti.read_text(encoding="utf-8")
        nr, tittel = kapittel_nr.get(filnavn, ("?", filnavn))

        # Kapitteltittel
        para = doc.add_heading(f"{nr} {tittel}", level=1)
        for run in para.runs:
            sett_font(run, størrelse=16, fet=True)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(10)

        # Fjern første H1-linje fra markdown (allerede lagt til over)
        linjer    = tekst.splitlines()
        uten_h1   = "\n".join(l for l in linjer if not l.startswith("# "))

        parse_markdown(doc, uten_h1)
        legg_til_sideskift(doc)

    # --- Referanser ---
    print("Behandler: referanser.md")
    para = doc.add_heading("Referanser", level=1)
    for run in para.runs:
        sett_font(run, størrelse=16, fet=True)

    ref_tekst = REFERANSER.read_text(encoding="utf-8")
    for linje in ref_tekst.splitlines():
        if not linje.strip() or linje.startswith("#"):
            continue
        if linje.startswith("---"):
            continue
        para = doc.add_paragraph()
        para.paragraph_format.left_indent       = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(-1.25)  # Hengende innrykk
        para.paragraph_format.space_after       = Pt(6)
        legg_til_løpetekst(para, linje)

    # Vedlegg med Python-kode er ikke del av rapporten

    # --- Lagre ---
    doc.save(OUT)
    print(f"\nLagret: {OUT}")
    print(f"Sider (estimat): ~{sum(1 for _ in doc.paragraphs) // 35 + 1}")


if __name__ == "__main__":
    main()
